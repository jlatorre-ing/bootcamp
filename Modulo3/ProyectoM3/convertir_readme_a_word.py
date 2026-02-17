from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Ruta del README
script_dir = os.path.dirname(os.path.abspath(__file__))
readme_path = os.path.join(script_dir, 'README.md')
output_path = os.path.join(script_dir, 'README.docx')

# Leer el README
with open(readme_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Crear documento Word
doc = Document()

# Procesar el contenido línea por línea
lines = content.split('\n')

for line in lines:
    if not line.strip():
        # Línea vacía
        doc.add_paragraph()
    elif line.startswith('# '):
        # Título nivel 1
        p = doc.add_heading(line.replace('# ', ''), level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        # Título nivel 2
        doc.add_heading(line.replace('## ', ''), level=2)
    elif line.startswith('### '):
        # Título nivel 3
        doc.add_heading(line.replace('### ', ''), level=3)
    elif line.startswith('```'):
        # Bloque de código (ignorar marcadores)
        continue
    elif line.startswith('- '):
        # Lista con viñeta
        doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
    elif line.startswith('| '):
        # Tabla (detectar filas de tabla)
        continue
    else:
        # Párrafo normal
        doc.add_paragraph(line)

# Guardar documento
doc.save(output_path)
print(f"✅ Archivo Word creado: {output_path}")
